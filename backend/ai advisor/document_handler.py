from document_cache import document_cache
from openrouter_client import openrouter_client
from config import MODEL_NAME
import logging
import os
from docx import Document as DocxDocument
import PyPDF2
from typing import Dict, Any, Optional, List
import json
import uuid

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

def extract_text_from_docx(file_path: str) -> str:
    """
    Extract text from a .docx file
    """
    try:
        doc = DocxDocument(file_path)
        full_text = []
        
        # Extract text from paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text)
                if row_text:
                    full_text.append(" | ".join(row_text))
        
        return "\n".join(full_text)
    except Exception as e:
        logger.error(f"Error extracting text from DOCX: {str(e)}")
        raise

def extract_text_from_pdf(file_path: str) -> str:
    """
    Extract text from a PDF file
    """
    try:
        text = []
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text.append(page.extract_text())
        return "\n".join(text)
    except Exception as e:
        logger.error(f"Error extracting text from PDF: {str(e)}")
        raise

def detect_document_type(filename: str, content: str = None) -> str:
    """
    Detect the type of document based on filename and content
    """
    filename = filename.lower()
    
    # Detect by filename
    if any(term in filename for term in ["resume", "cv", "curriculum"]):
        return "resume"
    elif any(term in filename for term in ["transcript", "academic", "grades"]):
        return "transcript"
    elif any(term in filename for term in ["book", "review", "literature"]):
        return "book_review"
    
    # Default document type
    return "general"

async def analyze_document_with_ai(doc_id: str, doc_type: str) -> str:
    """
    Send document to AI for analysis, using the cached segments
    
    Args:
        doc_id: The document ID in cache
        doc_type: Type of document (resume, transcript, etc.)
        
    Returns:
        Analysis text from AI
    """
    try:
        # Get metadata to know how many segments
        metadata = document_cache.get_document_metadata(doc_id)
        if not metadata:
            return "Document not found in cache."
            
        # Get all segments
        segments = []
        for i in range(metadata.get('segment_count', 0)):
            segment = document_cache.get_document_segment(doc_id, i)
            if segment:
                segments.append(segment)
                
        # Combine segments (with limits to prevent token size issues)
        max_content_length = 8000  # Adjust based on your model's context window
        content = "\n\n".join(segments)
        if len(content) > max_content_length:
            content = content[:max_content_length] + "... [Document truncated due to length]"
            
        # Customize prompt based on document type
        if doc_type.lower() == "resume":
            prompt = (
                "You are a career advisor. Analyze the following resume and provide: "
                "1. Strengths, 2. Weaknesses, 3. Suggestions for improvement, and "
                "4. 3 suitable career paths. Resume:\n\n"
                f"{content}\n\n"
                "Your analysis:"
            )
        elif doc_type.lower() == "transcript":
            prompt = (
                "You are a career counselor. Based on the following academic transcript, "
                "suggest 3 career paths and 2 skills to develop further. Transcript:\n\n"
                f"{content}\n\n"
                "Your suggestions:"
            )
        elif doc_type.lower() == "book_review":
            prompt = (
                "You are a literature expert. Summarize the following book review, highlighting "
                "key points, author's perspective, and main takeaways. Book review:\n\n"
                f"{content}\n\n"
                "Your summary:"
            )
        else:
            prompt = (
                "You are a document analysis expert. Analyze the following document, providing "
                "a detailed summary and key insights. Document:\n\n"
                f"{content}\n\n"
                "Your analysis:"
            )

        # Prepare messages for the AI
        messages = [
            {
                "role": "system", 
                "content": "You are a helpful document analyzer. Provide specific, clear analysis based on the provided content."
            },
            {
                "role": "user", 
                "content": prompt
            }
        ]

        logger.info(f"Sending {doc_type} document for analysis")
        completion = openrouter_client.chat.completions.create(
            messages=messages,
            model=MODEL_NAME,
            temperature=0.5,
            frequency_penalty=0.5,
            presence_penalty=0.5,
            max_tokens=1000
        )

        return completion.choices[0].message.content.strip()
        
    except Exception as e:
        logger.error(f"Error in analyze_document_with_ai: {str(e)}")
        return "Sorry, I encountered an error analyzing the document. Please try again."

async def process_document(file_path: str, file_name: str) -> Dict[str, Any]:
    """
    Process a document file, cache it, and return analysis
    
    Args:
        file_path: Path to the uploaded file
        file_name: Name of the uploaded file
        
    Returns:
        Dictionary with document analysis results and cache ID
    """
    try:
        # Get file extension
        file_ext = os.path.splitext(file_name)[1].lower()
        
        # Extract text based on file type
        if file_ext == '.docx':
            text_content = extract_text_from_docx(file_path)
        elif file_ext == '.pdf':
            text_content = extract_text_from_pdf(file_path)
        else:
            # For other file types, try to read as text
            with open(file_path, 'r', encoding='utf-8') as file:
                text_content = file.read()
        
        # Detect document type
        doc_type = detect_document_type(file_name, text_content)
        logger.info(f"Detected document type: {doc_type} for {file_name}")
        
        # Store in cache
        doc_id = document_cache.add_document(file_name, text_content, doc_type)
        
        # Analyze with AI
        analysis = await analyze_document_with_ai(doc_id, doc_type)
        
        return {
            "document_type": doc_type,
            "file_name": file_name,
            "content_length": len(text_content),
            "doc_id": doc_id,
            "analysis": analysis
        }
        
    except Exception as e:
        logger.error(f"Error processing document {file_name}: {str(e)}")
        return {
            "document_type": "unknown",
            "file_name": file_name,
            "error": str(e),
            "analysis": "I encountered an error while processing your document. Please try again or upload a different file format."
        }

def get_document_segments_for_context(doc_id: str, segment_indices: List[int] = None) -> str:
    """
    Get document segments for context inclusion
    
    Args:
        doc_id: Document ID in cache
        segment_indices: Specific segment indices to retrieve (None for summary)
        
    Returns:
        Document content for context
    """
    if not doc_id:
        return ""
        
    metadata = document_cache.get_document_metadata(doc_id)
    if not metadata:
        return ""
        
    if segment_indices:
        # Get specific segments
        segments = []
        for idx in segment_indices:
            segment = document_cache.get_document_segment(doc_id, idx)
            if segment:
                segments.append(segment)
        return "\n\n".join(segments)
    else:
        # Get summary (first few segments)
        return document_cache.get_document_summary(doc_id)