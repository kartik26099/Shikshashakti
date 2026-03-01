import os
from typing import List, BinaryIO
import PyPDF2
import docx
import re
from sentence_transformers import SentenceTransformer
from db_setup import SessionLocal, Document, DocumentChunk
from db_setup import Base, engine
from functools import lru_cache
import requests
from dotenv import load_dotenv

# Initialize database tables
Base.metadata.create_all(engine)

# Load environment variables
load_dotenv()

# OpenRouter setup for summarization
API_URL = os.getenv("LLM_API_URL", "https://openrouter.ai/api/v1/chat/completions")
API_KEY = os.getenv("LLM_API_KEY", "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")
MODEL_NAME = os.getenv("MODEL_NAME", "meta-llama/llama-3.1-8b-instruct:free")

HEADERS = {
    "Authorization": f"Bearer {API_KEY}",
    "Content-Type": "application/json",
    "HTTP-Referer": "http://localhost:8000",
    "X-Title": "AI Faculty Document Processor"
}

# Load embedding model (using sentence-transformers instead of OpenAI)
model = SentenceTransformer('all-MiniLM-L6-v2')
@lru_cache(maxsize=1000)
def extract_text_from_file(file: BinaryIO, filename: str) -> str:
    """Extract text from uploaded files based on file type."""
    if filename.endswith('.pdf'):
        pdf_reader = PyPDF2.PdfReader(file)
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
        return text
    elif filename.endswith('.docx'):
        doc = docx.Document(file)
        text = ""
        for para in doc.paragraphs:
            text += para.text + "\n"
        return text
    elif filename.endswith('.txt'):
        return file.read().decode('utf-8')
    else:
        raise ValueError(f"Unsupported file format: {filename}")

def split_text_into_chunks(text: str, chunk_size: int = 1000, overlap: int = 150) -> List[str]:
    """Split text into overlapping chunks with a smaller size."""
    print(f"Starting split_text_into_chunks with text of length {len(text)}")
    
    if not text:
        print("Text is empty, returning empty list")
        return []
    
    # For very small texts, just return the whole text as one chunk
    if len(text) <= chunk_size:
        print(f"Text length ({len(text)}) is smaller than chunk_size ({chunk_size}), returning whole text as one chunk")
        return [text]
    
    # Clean text
    print("Cleaning text...")
    text = re.sub(r'\s+', ' ', text).strip()
    print(f"Cleaned text length: {len(text)}")
    
    # Split into chunks
    chunks = []
    start = 0
    print(f"Starting chunking loop with start={start}, len(text)={len(text)}")
    
    while start < len(text):
        end = min(start + chunk_size, len(text))
        print(f"Iteration: start={start}, end={end}")
        
        # Find a good breaking point (space) if not at the end
        if end < len(text):
            print(f"Finding breaking point near position {end}")
            original_end = end
            while end > start and text[end] != ' ':
                end -= 1
                
            if end == start:  # No space found, just cut at chunk_size
                print(f"No space found between {start} and {original_end}, using original end")
                end = original_end
            else:
                print(f"Found space at position {end}")
        
        chunk = text[start:end]
        print(f"Adding chunk of length {len(chunk)}")
        chunks.append(chunk)
        
        new_start = end - overlap
        print(f"New start position: {new_start} (end={end}, overlap={overlap})")
        
        # Prevent infinite loop
        if new_start <= start:
            print(f"WARNING: New start position ({new_start}) <= current start ({start}), advancing to end")
            new_start = end
            
        start = new_start
        
        # Break if we've reached the end
        if end >= len(text):
            print("Reached end of text, breaking loop")
            break
    
    print(f"Finished chunking, created {len(chunks)} chunks")
    return chunks


def generate_embedding(text: str) -> List[float]:
    """Generate embedding vector for text using sentence-transformers with caching."""
    try:
        return model.encode(text).tolist()
    except Exception as e:
        print(f"Error generating embedding: {str(e)}")
        # Return a zero vector as fallback
        return [0.0] * 384  # Default dimension for 'all-MiniLM-L6-v2'

def process_document(file: BinaryIO, filename: str, title: str) -> int:
    """Process document, extract text, generate summary, and store chunks without using Pinecone."""
    try:
        print(f"Starting to process document: {filename}")
        print(f"Document title: {title}")
        
        # Extract text from file
        print("Extracting text from file...")
        text = extract_text_from_file(file, filename)
        print(f"Extracted text length: {len(text)} characters")
        
        # Generate document summary
        print("Generating document summary...")
        summary = generate_document_summary(text, title)
        print(f"Generated summary length: {len(summary)} characters")
        
        # Extract key topics from summary
        print("Extracting key topics...")
        topics = extract_key_topics_from_summary(summary)
        print(f"Extracted topics: {topics}")
        
        # Create document in database
        print("Creating document in database...")
        db = SessionLocal()
        try:
            import json
            document = Document(
                title=title, 
                content=text,
                summary=summary,
                topics=json.dumps(topics)
            )
            db.add(document)
            db.commit()
            db.refresh(document)
            print(f"Document created with ID: {document.id}")
            
            # Split text into chunks
            print(f"Splitting text into chunks (size={len(text)}, chunk_size=1000, overlap=150)...")
            chunks = split_text_into_chunks(text, chunk_size=1000, overlap=150)
            print(f"Created {len(chunks)} chunks")
            
            # Store full chunks in database
            print("Storing chunks in database...")
            store_full_chunks(document.id, chunks)
            print("Chunks stored successfully")
            
            return document.id
        finally:
            db.close()
    except Exception as e:
        print(f"ERROR in process_document: {str(e)}")
        import traceback
        traceback.print_exc()
        raise  # Re-raise the exception to be caught by the endpoint handler

def store_full_chunks(document_id: int, chunks: List[str]) -> None:
    """Store full text chunks in database for retrieval."""
    print(f"Storing {len(chunks)} chunks for document ID {document_id}")
    
    db = SessionLocal()
    try:
        # Delete any existing chunks for this document
        print(f"Deleting existing chunks for document ID {document_id}")
        db.query(DocumentChunk).filter(DocumentChunk.document_id == document_id).delete()
        
        # Insert new chunks
        print("Inserting new chunks...")
        for i, chunk in enumerate(chunks):
            print(f"Creating chunk record {i+1}/{len(chunks)} with length {len(chunk)}")
            chunk_record = DocumentChunk(
                document_id=document_id,
                chunk_index=i,
                content=chunk
            )
            db.add(chunk_record)
        
        print("Committing chunks to database...")
        db.commit()
        print("Chunks committed successfully")
    except Exception as e:
        db.rollback()
        print(f"ERROR storing document chunks: {str(e)}")
        import traceback
        traceback.print_exc()
        raise
    finally:
        db.close()

def cosine_similarity(vec1, vec2):
    """Calculate cosine similarity between two vectors."""
    dot_product = sum(a*b for a, b in zip(vec1, vec2))
    magnitude1 = sum(a*a for a in vec1) ** 0.5
    magnitude2 = sum(b*b for b in vec2) ** 0.5
    if magnitude1 * magnitude2 == 0:
        return 0
    return dot_product / (magnitude1 * magnitude2)

def generate_document_summary(document_content: str, document_title: str) -> str:
    """Generate a comprehensive summary of the document for better RAG performance."""
    
    # Truncate content if too long (keep first 8000 characters for summary)
    content_for_summary = document_content[:8000]
    
    system_prompt = """You are an expert at creating comprehensive document summaries for educational purposes.
    
    Your task is to create a detailed summary that includes:
    1. Main topics and themes covered
    2. Key concepts and definitions
    3. Important facts and data points
    4. Structure and organization of the content
    5. Learning objectives and takeaways
    
    The summary should be well-structured and comprehensive enough to answer questions about the document content.
    """
    
    user_prompt = f"""Please create a comprehensive summary of this document:
    
    Title: {document_title}
    
    Content:
    {content_for_summary}
    
    Create a detailed summary that captures all important information and can be used to answer questions about the document content."""
    
    try:
        data = {
            "model": MODEL_NAME,
            "messages": [
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            "temperature": 0.3,
            "max_tokens": 1500
        }
        
        response = requests.post(API_URL, headers=HEADERS, json=data)
        response.raise_for_status()
        result = response.json()
        
        summary = result["choices"][0]["message"]["content"].strip()
        return summary
        
    except Exception as e:
        print(f"Error generating document summary: {str(e)}")
        # Fallback: create a basic summary
        return f"Document Summary for '{document_title}': This document contains educational content covering various topics. The content has been processed and is available for questions and quiz generation."

def extract_key_topics_from_summary(summary: str) -> List[str]:
    """Extract key topics from the document summary for better categorization."""
    
    system_prompt = """You are an expert at identifying key educational topics from document summaries.
    Extract 3-5 main topics that would be useful for categorization and quiz generation.
    Return only a JSON array of topic strings."""
    
    user_prompt = f"""From this document summary, extract 3-5 key educational topics:
    
    {summary}
    
    Return only a JSON array of topic strings. Example: ["Topic 1", "Topic 2", "Topic 3"]"""
    
    try:
        data = {
            "model": MODEL_NAME,
            "messages": [
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            "temperature": 0.7,
            "max_tokens": 300,
            "response_format": {"type": "json_object"}
        }
        
        response = requests.post(API_URL, headers=HEADERS, json=data)
        response.raise_for_status()
        result = response.json()
        
        response_text = result["choices"][0]["message"]["content"]
        
        # Parse JSON response
        import json
        try:
            topics_data = json.loads(response_text)
            if isinstance(topics_data, list):
                return topics_data
            elif isinstance(topics_data, dict) and "topics" in topics_data:
                return topics_data["topics"]
            else:
                return ["General Knowledge"]
        except json.JSONDecodeError:
            # Fallback: extract topics using regex
            topics = re.findall(r'"([^"]+)"', response_text)
            return topics if topics else ["General Knowledge"]
            
    except Exception as e:
        print(f"Error extracting topics: {str(e)}")
        return ["General Knowledge"]